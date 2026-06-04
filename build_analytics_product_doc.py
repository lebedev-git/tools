from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).with_name("analytics_product_architecture_2026.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = RGBColor(30, 36, 44)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Heading 1", 15, RGBColor(16, 48, 82)),
        ("Heading 2", 12, RGBColor(24, 59, 94)),
        ("Heading 3", 10.5, RGBColor(40, 72, 104)),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(9.2)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.left_indent = Cm(0.42)

    compact = doc.styles.add_style("Compact Table", WD_STYLE_TYPE.TABLE)
    compact.font.name = "Aptos"
    compact.font.size = Pt(8.5)
    callout = doc.styles.add_style("Callout Table", WD_STYLE_TYPE.TABLE)
    callout.font.name = "Aptos"
    callout.font.size = Pt(9)


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Архитектура продукта: аналитика и протоколы"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(102, 112, 128)

    footer = section.footer.paragraphs[0]
    footer.text = "lebedev-git/tools · проектная спецификация 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(102, 112, 128)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Compact Table"
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, text in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E8EEF6")
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(24, 36, 52)

    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    return table


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Callout Table"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F7FB")
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(24, 59, 94)
    run.font.size = Pt(9.5)
    paragraph.add_run("\n" + body)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_process_detail(
    doc: Document,
    title: str,
    goal: str,
    user_flow: str,
    system_flow: str,
    saved_data: str,
    ui_state: str,
    failures: str,
    success: str,
) -> None:
    doc.add_heading(title, level=2)
    add_table(
        doc,
        ["Блок", "Описание"],
        [
            ["Цель", goal],
            ["Действия пользователя", user_flow],
            ["Работа системы", system_flow],
            ["Что сохраняется", saved_data],
            ["Что видно в UI", ui_state],
            ["Ошибки и retry", failures],
            ["Успешный результат", success],
        ],
        widths=[3.2, 13.3],
    )


def build() -> None:
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Новая архитектура продукта 2026")
    run.font.name = "Aptos Display"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = RGBColor(14, 45, 76)
    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Два соседних инструмента в одном продукте: Analytics Tool и Protocol Tool. "
        "Общая инфраструктура, но независимая бизнес-логика, независимые шаблоны и независимые документы."
    )
    subtitle.runs[0].font.size = Pt(10.5)
    subtitle.runs[0].font.color.rgb = RGBColor(74, 85, 101)

    add_callout(
        doc,
        "Ключевое решение",
        "Аналитика и протоколы не объединяются в один workflow. Они находятся рядом в UI shell, используют общий backend, очередь, LLM adapter, Outline integration и дизайн-систему, но каждый инструмент имеет собственные шаблоны, источники, статусы документа и пользовательские сценарии.",
    )

    doc.add_heading("1. Текущий контекст", level=1)
    add_bullets(
        doc,
        [
            "Текущий UI уже показывает рядом два раздела: «Аналитика» и «Протоколы». Это нужно сохранить как продуктовую модель.",
            "Для аналитики выгружен n8n workflow analytics-note: Яндекс Формы -> промпты -> LLM -> медиа/изображение -> Outline.",
            "Для настроек промптов выгружен workflow analytics-prompt-settings. Логику нужно перенести в код и таблицы версий.",
            "Выгружено 7 n8n workflows; часть относится к протоколам или backup/demo. Новый продукт не должен развивать n8n, а должен использовать exports как legacy reference.",
            "GitHub-репозиторий lebedev-git/tools пустой, поэтому этот документ является первым архитектурным артефактом для дальнейшей реализации.",
        ],
    )

    doc.add_heading("2. Модель продукта", level=1)
    add_table(
        doc,
        ["Часть продукта", "Назначение", "Связь с другими частями"],
        [
            ["Product Shell", "Общая оболочка: навигация, авторизация, layout, дизайн-система, глобальные настройки.", "Содержит оба инструмента, но не смешивает их сценарии."],
            ["Analytics Tool", "Аналитические записки: День 1, День 2, общая аналитика, NPS, фото/медиа, публикация.", "Не использует протокол как источник и не создает протоколы."],
            ["Protocol Tool", "Создание, структурирование, редактирование, согласование и публикация протоколов.", "Не запускает аналитику и не входит в конструктор аналитики."],
            ["Shared Platform", "База, очередь, LLM adapter, Outline integration, prompt versions, audit log, provider requests.", "Обслуживает оба инструмента одинаковыми техническими правилами."],
        ],
        widths=[3.0, 7.0, 6.5],
    )
    add_callout(
        doc,
        "Жесткое правило границы",
        "Нельзя проектировать общий бизнес-процесс для двух инструментов. Если в будущем понадобится ручная ссылка между документами, она должна быть обычной ссылкой в Documents, а не бизнес-зависимостью между инструментами.",
    )

    doc.add_heading("3. Выбор стека", level=1)
    add_table(
        doc,
        ["Задача", "Варианты", "Решение", "Почему"],
        [
            ["Full-stack", "Next.js; TanStack Start; React Router/Remix", "Next.js App Router", "Один продукт с UI, API routes, server actions и обычным Node deployment без Docker."],
            ["Очереди", "Graphile Worker; BullMQ; Trigger.dev/Inngest/Temporal", "Graphile Worker", "PostgreSQL-only очередь без Redis и внешней платформы; подходит для локального запуска и сервера."],
            ["База/ORM", "Drizzle; Prisma; Kysely", "Drizzle ORM", "Типизация, прозрачные SQL-миграции, небольшой runtime, удобная работа с явными таблицами аудита."],
            ["UI kit", "shadcn/ui; Mantine; HeroUI/Chakra", "shadcn/ui", "Компоненты живут в коде, хорошо подходят для компактной операторской консоли и редактора документов."],
            ["Визуализация процессов", "React Flow/xyflow; Rete.js; Craft.js", "React Flow/xyflow", "Подходит для графа шагов аналитики и visual progress, но не превращает продукт в n8n-клон."],
            ["LLM", "Vercel AI SDK; прямой OpenAI-compatible client; LangChain.js", "Provider adapter + Vercel AI SDK при необходимости", "Позволяет менять провайдера и фиксировать каждый запрос в provider_requests."],
        ],
        widths=[2.7, 4.0, 3.4, 6.4],
    )

    doc.add_heading("4. Целевая архитектура", level=1)
    add_bullets(
        doc,
        [
            "apps/web: Next.js UI shell, страницы Analytics Tool и Protocol Tool, API routes, server actions, auth.",
            "apps/worker: Graphile Worker tasks для обоих инструментов: analytics jobs и protocol jobs живут отдельно, но используют один worker runtime.",
            "packages/core: типы процессов, state machine, валидаторы, общие event/status helpers.",
            "packages/analytics: шаблоны аналитики, нормализация форм, расчет метрик, сборка prompt payloads, публикация аналитических документов.",
            "packages/protocols: шаблоны протоколов, парсинг источников, извлечение решений/задач, workflow согласования, публикация протоколов.",
            "packages/integrations: Yandex Forms, Outline, LLM provider, file storage, image/media service.",
            "packages/db: schema, migrations, repositories.",
        ],
    )
    add_table(
        doc,
        ["Shared Platform", "Что общее для Analytics и Protocol"],
        [
            ["UI shell", "Одинаковая навигация, layout, дизайн-система, responsive правила, auth state."],
            ["Database", "Общие таблицы запусков, событий, provider requests, prompt settings и публикаций."],
            ["Queue", "Один механизм постановки фоновых задач, но разные task handlers для аналитики и протоколов."],
            ["LLM adapter", "Единый слой провайдера, лимитов, structured output, retry и логирования стоимости/latency."],
            ["Outline integration", "Единый publish/update/version adapter, разные document templates."],
            ["Prompt settings", "Общая система версий, но разные группы промптов: analytics_* и protocol_*."],
        ],
        widths=[3.5, 13.0],
    )

    doc.add_heading("5. Модель данных и индексация", level=1)
    add_table(
        doc,
        ["Таблица", "Назначение"],
        [
            ["process_runs", "Один запуск любого инструмента. Поля: tool_type = analytics | protocol, status, created_by, input_summary, result_document_id."],
            ["process_run_steps", "Шаги запуска. Для аналитики и протоколов разные step_type, но одинаковые статусы и retry rules."],
            ["step_events", "Лента событий: validation, fetch, normalize, llm, publish, error, retry, user_action."],
            ["provider_requests", "Все запросы к LLM, image service, Outline и внешним API с latency, tokens, request_hash, response_ref."],
            ["prompt_settings", "Версии промптов по namespace: analytics.day1, analytics.day2, analytics.nps, protocol.meeting, protocol.session."],
            ["outline_publications", "Факты публикации: document_type, outline_id, url, version, published_at, source_run_id."],
            ["document_versions", "Версии черновиков и опубликованных документов, чтобы retry не затирал утвержденные результаты."],
            ["form_snapshots", "Snapshots ответов форм для аналитики и протоколов, если протокол создается из формы."],
            ["protocol_records", "Основная карточка протокола: title, date, participants, status, template_id, current_version_id."],
            ["protocol_action_items", "Задачи из протокола: text, owner, due_date, status, source_section_id."],
            ["protocol_decisions", "Зафиксированные решения: text, context, decision_date, responsible_party."],
            ["analytics_documents", "Итоговые аналитические записки и связи с выбранными сессиями/днями."],
        ],
        widths=[4.0, 12.5],
    )
    add_callout(
        doc,
        "Воспроизводимость",
        "Каждый результат должен ссылаться на конкретные versions: шаблон процесса, prompt, источники, snapshot данных и provider requests. Это одинаково важно и для аналитики, и для протоколов.",
    )

    doc.add_heading("6. Analytics Tool: конструктор и логика", level=1)
    doc.add_paragraph(
        "Конструктор аналитики работает только внутри Analytics Tool. Протоколы не появляются в конструкторе аналитики, и он не предлагает блоки протокола. Его задача — собрать аналитический сценарий из ограниченного каталога блоков и показать live-граф выполнения."
    )
    add_table(
        doc,
        ["Блок", "Настройки пользователя", "Результат"],
        [
            ["День 1", "Дата/сессия, входная форма, выходная форма, prompt preset, модель.", "Аналитическая секция или документ по первому дню."],
            ["День 2", "Включить/выключить, форма второго дня, режим сравнения с Днем 1.", "Аналитика второго дня и сравнительные выводы."],
            ["Общая аналитика", "Выбранные дни, режим синтеза, формат рекомендаций.", "Сводная записка по нескольким источникам."],
            ["NPS", "Источник оценок, шкала, сегменты, включить в документ или отдельный отчет.", "Расчет NPS, распределение и текстовая интерпретация."],
            ["Фото/медиа", "Загружать фото: да/нет; лимит; подписи; использовать для изображения/контекста.", "Медиа-блоки или dashboard image."],
            ["Публикация", "Outline collection/folder, режим create/update, название документа.", "Ссылка на опубликованный документ и версия публикации."],
        ],
        widths=[3.2, 7.2, 6.0],
    )

    add_process_detail(
        doc,
        "6.1 Аналитика День 1",
        "Сформировать аналитическую записку по первой части сессии на основе входных и выходных форм.",
        "Оператор открывает «Аналитика», выбирает шаблон «День 1», выбирает дату/сессию, проверяет найденные формы и нажимает Run.",
        "API валидирует параметры, создает process_run, worker загружает формы, нормализует ответы, считает базовые метрики, строит prompt payload, вызывает LLM, валидирует структуру ответа, публикует документ в Outline.",
        "process_run, steps, step_events, form_snapshots, provider_requests, analytics_documents, outline_publications.",
        "Таблица сессий, граф шагов, прогресс normalize/LLM/publish, ссылка на Outline, drawer с логами.",
        "Если форма недоступна — step failed и предлагается retry. Если LLM нарушил структуру — автоматический retry с corrective prompt. Если Outline недоступен — можно повторить только publish step.",
        "Документ Outline создан, ссылка сохранена, все шаги succeeded, snapshot и версия prompt зафиксированы.",
    )
    add_process_detail(
        doc,
        "6.2 Аналитика День 2",
        "Добавить аналитику второго дня как отдельный аналитический блок, опционально сравнить с Днем 1.",
        "Оператор в конструкторе включает блок «День 2», выбирает источник второго дня и режим: самостоятельный отчет или сравнение с Днем 1.",
        "Worker загружает day2 source, нормализует ответы, при включенном сравнении подтягивает snapshot Дня 1, строит prompt, получает выводы, публикует отдельный документ или добавляет секцию.",
        "form_snapshots day2, analytics_documents, связи между day1/day2 snapshots внутри analytics metadata.",
        "В графе появляется отдельная ветка Day 2. UI показывает, есть ли зависимость от Day 1 и какие данные используются для сравнения.",
        "Если Day 1 нужен для сравнения, но не выбран, блок blocked до исправления настроек. При ошибке day2 формы retry не трогает уже готовый Day 1.",
        "Аналитика второго дня создана, связь с выбранным режимом сохранена, пользователь видит итоговую ссылку.",
    )
    add_process_detail(
        doc,
        "6.3 Общая аналитика",
        "Собрать итоговый синтез по выбранным аналитическим источникам без участия Protocol Tool.",
        "Оператор выбирает «Общая аналитика», отмечает нужные дни/сессии, выбирает формат вывода и запускает процесс.",
        "Система загружает snapshots и/или готовые analytics sections, агрегирует метрики, строит synthesis prompt, валидирует результат и публикует итоговую записку.",
        "Агрегированный input payload, provider_requests, analytics_documents, outline_publications.",
        "UI показывает список включенных источников, предупреждения по неполным данным, итоговый progress и ссылку на документ.",
        "Если один из источников неполный, UI предлагает исключить источник или остановить запуск. Retry общей аналитики не пересоздает исходные документы.",
        "Создана итоговая аналитическая записка с воспроизводимым списком источников.",
    )
    add_process_detail(
        doc,
        "6.4 NPS",
        "Посчитать NPS и добавить интерпретацию как часть аналитического результата или отдельный отчет.",
        "Оператор включает блок NPS, выбирает источник оценок, шкалу и сегменты.",
        "Worker загружает ответы, считает promoters/passives/detractors, строит распределение, вызывает LLM для краткой интерпретации, сохраняет результат.",
        "nps_results внутри analytics metadata, provider_requests для интерпретации, document section/output.",
        "UI показывает расчет, распределение, сегменты, статус включения в документ.",
        "Если поле оценки не найдено, блок blocked и требует mapping. Если оценок мало, UI показывает предупреждение, но запуск можно продолжить.",
        "NPS рассчитан, интерпретация сохранена, документ обновлен или отдельный отчет опубликован.",
    )
    add_process_detail(
        doc,
        "6.5 Фото и медиа",
        "Опционально включить фото/медиа в аналитический сценарий без обязательности для протоколов.",
        "Оператор включает toggle «Загружать фото», добавляет файлы или выбирает источник, задает подписи и режим использования.",
        "Система валидирует типы и лимиты, сохраняет metadata, прикрепляет медиа к analytics run, при необходимости вызывает image/dashboard service.",
        "file metadata, media attachments, provider_requests image service, связи с analytics_document.",
        "UI показывает список файлов, статус обработки, предупреждения по размеру/типу, preview при наличии.",
        "Невалидный файл исключается до запуска. Ошибка image generation не должна ломать текстовую аналитику, если блок помечен optional.",
        "Медиа прикреплены к документу или использованы для dashboard image, результат проиндексирован.",
    )
    add_process_detail(
        doc,
        "6.6 Prompt settings и retry/re-run",
        "Управлять версиями промптов и безопасно повторять запуски без потери истории.",
        "Оператор редактирует prompt как draft, тестирует на sample-сессии, активирует версию. Для ошибки выбирает retry step или full re-run.",
        "Система сохраняет prompt versions, привязывает run к конкретной версии, retry перезапускает failed step и зависимые шаги, full re-run создает новый process_run.",
        "prompt_settings, process_runs, process_run_steps, provider_requests, document_versions.",
        "UI показывает активную версию prompt, историю запусков, отличия re-run от retry, ссылки на версии документов.",
        "Published документы не затираются без явного update/create new version. Idempotency key предотвращает случайные дубли.",
        "Промпты версионированы, повторный запуск прозрачен, старые результаты доступны для сравнения.",
    )

    doc.add_heading("7. Protocol Tool: отдельный инструмент", level=1)
    doc.add_paragraph(
        "Protocol Tool — самостоятельный инструмент рядом с аналитикой. Он не является источником аналитики, не входит в конструктор аналитики и не запускает аналитические сценарии. Его задача — подготовить структурированный протокол и опубликовать его в Outline."
    )
    add_table(
        doc,
        ["Элемент", "Описание"],
        [
            ["Источники", "Ручной ввод, файл, transcript, форма, текстовая заметка, существующий черновик."],
            ["Поля протокола", "Название, дата, участники, тема, повестка, тезисы, решения, задачи, ответственные, сроки, риски, приложения."],
            ["Статусы", "draft, generated, review, approved, published, archived."],
            ["Шаблоны", "Встреча, сессия, рабочее совещание, проектный протокол, кастомный шаблон."],
            ["Публикация", "Outline create/update с versioning. Approved/published версия не затирается повторной генерацией."],
        ],
        widths=[3.2, 13.3],
    )
    add_process_detail(
        doc,
        "7.1 Создание протокола",
        "Быстро собрать структурированный протокол по встрече, сессии или мероприятию.",
        "Пользователь открывает «Протоколы», нажимает «Новый протокол», выбирает тип протокола и источник данных.",
        "API создает protocol_record и process_run tool_type=protocol. Worker извлекает текст/структуру, запускает protocol prompt, формирует черновик секций, решений и задач.",
        "protocol_record, protocol_sections, protocol_decisions, protocol_action_items, document_versions, provider_requests.",
        "Список протоколов, мастер создания, редактор структуры, статус generated, правый inspector секции.",
        "Если источник не читается, протокол остается draft. Если LLM не выделил обязательные поля, UI показывает validation issues и позволяет ручное заполнение.",
        "Черновик протокола создан, структурные поля заполнены, документ готов к review.",
    )
    add_process_detail(
        doc,
        "7.2 Редактирование и согласование",
        "Дать пользователю контроль над протоколом до публикации.",
        "Пользователь редактирует секции, решения, задачи, ответственных и сроки; переводит документ в review или approved.",
        "Система сохраняет document_versions и структурные записи. При изменении задач/решений обновляет protocol_action_items и protocol_decisions.",
        "document_versions, protocol_* tables, step_events user_action.",
        "Редактор показывает статус, diff версий, список задач и решений, предупреждения по пустым обязательным полям.",
        "Нельзя перевести в approved, если нет обязательных полей. Ошибки сохранения показываются рядом с редактируемым блоком.",
        "Протокол согласован, имеет approved версию и готов к публикации.",
    )
    add_process_detail(
        doc,
        "7.3 Публикация протокола",
        "Опубликовать approved или generated протокол в Outline без связи с Analytics Tool.",
        "Пользователь выбирает папку/collection Outline и нажимает Publish.",
        "Worker собирает финальный markdown/html payload, вызывает Outline adapter, сохраняет outline_publication и переводит статус в published.",
        "outline_publications, document_versions, protocol_record.status, provider_requests Outline.",
        "UI показывает progress publish, ссылку Outline, версию публикации и дату.",
        "Если Outline недоступен, retry повторяет только publish step. Старый published документ не затирается без выбора update existing.",
        "Протокол опубликован, ссылка доступна в списке протоколов и разделе Документы.",
    )
    add_process_detail(
        doc,
        "7.4 Повторная генерация протокола",
        "Позволить улучшить draft/generated протокол, не ломая утвержденные версии.",
        "Пользователь нажимает Regenerate для draft/generated или Create new version для approved/published.",
        "Система создает новую document_version. Для approved/published сначала создает новую draft-версию и не меняет текущую опубликованную.",
        "document_versions, step_events, provider_requests, protocol_record.current_version_id.",
        "UI четко показывает текущую опубликованную версию и новую черновую версию.",
        "Запрещено silent overwrite для approved/published. При конфликте версий пользователь выбирает, какую версию сделать текущей.",
        "Новая версия создана безопасно, история сохранена, published версия не потеряна.",
    )

    doc.add_heading("8. State machine и выполнение", level=1)
    add_table(
        doc,
        ["Тип", "Статусы", "Правила"],
        [
            ["process_run", "draft, validated, queued, running, waiting_external, succeeded, failed, cancelled", "Одинаковая модель для analytics и protocol; различается только tool_type и набор step_type."],
            ["process_run_step", "pending, blocked, queued, running, succeeded, failed, skipped, retrying", "Retry запускает failed step и зависимые шаги; независимые succeeded steps не пересоздаются."],
            ["analytics document", "draft, generated, published, superseded", "Аналитика может быть перезапущена как новый run; старый результат остается для сравнения."],
            ["protocol document", "draft, generated, review, approved, published, archived", "Approved/published нельзя затирать регенерацией; нужна новая версия."],
        ],
        widths=[3.5, 5.8, 7.2],
    )

    doc.add_heading("9. UI правила", level=1)
    add_bullets(
        doc,
        [
            "Left rail: Аналитика, Протоколы, Запуски, Документы, Настройки, Интеграции.",
            "Аналитика выглядит как операторская консоль: таблица сессий, конструктор блоков, граф выполнения, timeline/log, ссылки на отчеты.",
            "Протоколы выглядят как инструмент подготовки документа: список протоколов, создание нового, выбор шаблона, редактор структуры, панель решений/задач, публикация.",
            "Общий раздел «Запуски» показывает оба типа процессов с фильтром analytics | protocol.",
            "Общий раздел «Документы» показывает опубликованные документы, но сохраняет тип: analytics_note, nps_report, protocol.",
            "Конструктор аналитики не показывает протоколы. Protocol Tool не показывает аналитические блоки.",
            "Дизайн: компактность, 8px radius max, lucide icons, tooltips, стабильные размеры controls, доступный contrast, keyboard navigation.",
        ],
    )
    add_table(
        doc,
        ["Экран", "Analytics Tool", "Protocol Tool"],
        [
            ["Главный список", "Сессии, даты, количество входных/выходных ответов, быстрый Run.", "Протоколы, статус draft/review/published, дата, участники, последняя версия."],
            ["Создание", "Конструктор аналитического сценария из блоков Day 1/Day 2/NPS/Media.", "Wizard нового протокола: тип, источник, шаблон, первичная структура."],
            ["Рабочая область", "Граф шагов и таблицы данных.", "Редактор структуры документа и панели задач/решений."],
            ["Результат", "Аналитическая записка или NPS отчет.", "Протокол в Outline."],
        ],
        widths=[3.2, 6.7, 6.6],
    )

    doc.add_heading("10. Локальный запуск без Docker", level=1)
    add_numbered(
        doc,
        [
            "Установить Node.js 22+, pnpm, PostgreSQL 16/17 локально.",
            "Создать базу product_tools и пользователя с минимальными правами.",
            "Заполнить .env.local: DATABASE_URL, OUTLINE_API_URL, OUTLINE_TOKEN, YANDEX_FORMS_TOKEN, LLM_BASE_URL, LLM_API_KEY, STORAGE_PATH.",
            "Запустить pnpm install, pnpm db:migrate, pnpm dev и pnpm worker:dev.",
            "На сервере запускать web и worker через systemd/PM2, использовать Caddy/Nginx для TLS и reverse proxy.",
        ],
    )

    doc.add_heading("11. Миграция из n8n", level=1)
    add_numbered(
        doc,
        [
            "Зафиксировать все n8n exports как read-only legacy reference.",
            "Перенести analytics-note в typed analytics steps: fetch forms, normalize, metrics, prompt, LLM, media, Outline publish.",
            "Перенести prompt settings в таблицу prompt_settings с namespace analytics_* и protocol_*.",
            "Разобрать workflow протоколов отдельно и перенести его в Protocol Tool без связи с Analytics Tool.",
            "Собрать MVP Analytics: День 1 -> запуск -> progress -> Outline document.",
            "Собрать MVP Protocol: ручной ввод/файл -> generated draft -> review -> Outline publish.",
            "Сравнить 5-10 старых запусков n8n с новой аналитикой и отдельно проверить 3-5 протоколов.",
            "После проверки отключить n8n webhooks и перенести public endpoints на новый backend.",
        ],
    )

    doc.add_heading("12. Риски и решения", level=1)
    add_table(
        doc,
        ["Риск", "Решение"],
        [
            ["Смешение аналитики и протоколов", "Закрепить tool_type, отдельные страницы, отдельные templates и запрет protocol blocks в analytics constructor."],
            ["LLM возвращает нестабильный формат", "Structured output schema, validation, corrective retry, raw response archive."],
            ["Повторный запуск дублирует документы", "Idempotency key и document_versions; retry не равен full re-run."],
            ["Опубликованный протокол перезаписан", "Approved/published protocol только через new version, без silent overwrite."],
            ["Формы меняют структуру", "Form schema snapshots и явные mapping rules."],
            ["No Docker усложняет сервер", "systemd units, healthcheck endpoint, backup scripts, journald logs, миграции."],
        ],
        widths=[4.6, 11.9],
    )

    doc.add_heading("13. Исследованные ориентиры", level=1)
    add_table(
        doc,
        ["Категория", "Источники"],
        [
            ["Full-stack", "Next.js docs/GitHub; TanStack Start/Router GitHub; React Router GitHub."],
            ["Очереди", "Graphile Worker GitHub/docs; BullMQ GitHub/docs; Trigger.dev GitHub/docs; Inngest; Temporal TypeScript SDK."],
            ["UI", "shadcn/ui GitHub; Mantine GitHub; HeroUI/Chakra; Reddit r/reactjs обсуждение shadcn/ui 2026."],
            ["Конструктор", "xyflow/React Flow GitHub; Rete.js GitHub; Craft.js GitHub."],
            ["n8n alternatives", "Reddit r/n8n: open-source alternatives; Reddit про AI-native альтернативы n8n; статьи 2026 по Activepieces/Windmill/Make/Zapier."],
        ],
        widths=[3.2, 13.3],
    )
    doc.add_paragraph("Ключевые URL:")
    add_bullets(
        doc,
        [
            "https://github.com/vercel/next.js",
            "https://github.com/graphile/worker",
            "https://github.com/taskforcesh/bullmq",
            "https://github.com/triggerdotdev/trigger.dev",
            "https://github.com/inngest/inngest",
            "https://github.com/temporalio/sdk-typescript",
            "https://github.com/shadcn-ui/ui",
            "https://github.com/xyflow/xyflow",
            "https://github.com/retejs/rete",
            "https://github.com/prevwong/craft.js",
            "https://www.reddit.com/r/n8n/comments/1tmc065/looking_for_an_open_source_alternative_to_n8n/",
            "https://www.reddit.com/r/reactjs/comments/1tr6hoj/is_shadcnui_actually_worth_learning_in_2026_or/",
        ],
    )

    doc.add_heading("14. Следующий инженерный шаг", level=1)
    add_bullets(
        doc,
        [
            "Создать monorepo skeleton: apps/web, apps/worker, packages/db, packages/core, packages/analytics, packages/protocols, packages/integrations.",
            "Описать schema для process_runs, process_run_steps, prompt_settings, document_versions, outline_publications, protocol_* и analytics_documents.",
            "Сделать вертикальный срез Analytics Tool: День 1, progress graph, Outline document.",
            "Сделать вертикальный срез Protocol Tool: новый протокол, generated draft, review, publish.",
            "Провести UI polish pass: компактность, состояния, accessibility, mobile behavior, визуальная консистентность двух инструментов.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
